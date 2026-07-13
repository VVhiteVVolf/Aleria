from __future__ import annotations

from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from language_data import ALPHABET, PREFIXES, SUFFIXES, MALE_NAMES, FEMALE_NAMES, LEXICON


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parent
OUT = ROOT / "output" / "Kanaanith_Sprach_und_Namensbibel_1.0.docx"
CHART = WORKSPACE / "kanaanith-font" / "dist" / "Kanaanith_Zeichentafel_1.000.png"
SPECIMEN = WORKSPACE / "kanaanith-font" / "dist" / "Kanaanith_Schriftprobe_1.000.png"

INK = "34271C"
LAPIS = "1C596D"
GOLD = "B3883B"
PAPYRUS = "F0E5C8"
PALE = "F8F2E3"
RULE = "C6AD79"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd"); tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=65, start=105, bottom=65, end=105):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for name, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        el = tc_mar.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}"); tc_mar.append(el)
        el.set(qn("w:w"),str(value)); el.set(qn("w:type"),"dxa")


def set_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
    tc_w.set(qn("w:w"),str(int(inches*1440))); tc_w.set(qn("w:type"),"dxa")


def repeat_header(row):
    pr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"),"true"); pr.append(el)


def no_split(row):
    pr = row._tr.get_or_add_trPr(); pr.append(OxmlElement("w:cantSplit"))


def borders(table):
    pr = table._tbl.tblPr
    bd = pr.first_child_found_in("w:tblBorders")
    if bd is None:
        bd = OxmlElement("w:tblBorders"); pr.append(bd)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4"); el.set(qn("w:color"),RULE); bd.append(el)


def field(paragraph, instruction):
    run = paragraph.add_run(); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    text=OxmlElement("w:instrText"); text.set(qn("xml:space"),"preserve"); text.text=instruction
    end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    run._r.extend([begin,text,end])


def styles(doc):
    normal=doc.styles["Normal"]; normal.font.name="DejaVu Sans"; normal.font.size=Pt(10.3); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.16
    for name,size,color,before,after in (("Title",28,INK,0,10),("Subtitle",13,GOLD,0,8),("Heading 1",17,LAPIS,18,8),("Heading 2",13.5,GOLD,13,6),("Heading 3",11.5,LAPIS,9,4)):
        s=doc.styles[name]; s.font.name="DejaVu Sans"; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=name!="Subtitle"
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for name,size,bold,color in (("Table Text",7.9,False,INK),("Table Head",8.0,True,WHITE)):
        if name not in doc.styles:
            from docx.enum.style import WD_STYLE_TYPE
            s=doc.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
        else: s=doc.styles[name]
        s.font.name="DejaVu Sans"; s.font.size=Pt(size); s.font.bold=bold; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_after=Pt(0); s.paragraph_format.line_spacing=1.0


def setup(doc):
    sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75); sec.header_distance=Inches(.33); sec.footer_distance=Inches(.33)
    sec.different_first_page_header_footer=True
    sec.first_page_header.paragraphs[0].text=""; sec.first_page_footer.paragraphs[0].text=""
    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=hp.add_run("KANA’ANITH · SPRACHBIBEL 1.0"); r.font.name="DejaVu Sans"; r.font.size=Pt(7.5); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(GOLD)
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("Südost-Tirnara  ·  "); r.font.size=Pt(7.5); r.font.color.rgb=RGBColor.from_string(LAPIS); field(fp,"PAGE")


def rule(doc):
    p=doc.add_paragraph(); pr=p._p.get_or_add_pPr(); bdr=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
    bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"14"); bottom.set(qn("w:space"),"1"); bottom.set(qn("w:color"),GOLD); bdr.append(bottom); pr.append(bdr)


def table(doc, headers, rows, widths, first_shaded=False):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; borders(t)
    hr=t.rows[0]; repeat_header(hr)
    for i,h in enumerate(headers):
        set_width(hr.cells[i],widths[i]); shade(hr.cells[i],LAPIS); margins(hr.cells[i]); p=hr.cells[i].paragraphs[0]; p.style="Table Head"; p.add_run(str(h))
    for ri,values in enumerate(rows):
        row=t.add_row(); no_split(row)
        for ci,value in enumerate(values):
            c=row.cells[ci]; set_width(c,widths[ci]); margins(c)
            if ri%2: shade(c,PALE)
            if first_shaded and ci==0: shade(c,PAPYRUS)
            p=c.paragraphs[0]; p.style="Table Text"; p.add_run(str(value)); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t


def bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.16); p.paragraph_format.space_after=Pt(3); p.add_run(item)


def cover(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(70); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("KANA’ANITH"); r.font.name="DejaVu Sans"; r.font.size=Pt(34); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(INK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Sprach-, Namens- und Schriftbibel Südost-Tirnaras"); r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(GOLD)
    rule(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18)
    r=p.add_run("22 Bedeutungszeichen · 200 Personennamen · 100 Kulturwörter · Monumentalfont"); r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(LAPIS)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30); p.paragraph_format.left_indent=Inches(.65); p.paragraph_format.right_indent=Inches(.65)
    r=p.add_run("Eine eigenständige Abjadsprache mit kanaanitischer Konsonantenlogik, ägyptisch-monumentaler Bildkraft und den Handels-, Tempel- und Flusskulturen Tirnaras."); r.font.name="DejaVu Serif"; r.font.size=Pt(11.5); r.font.italic=True; r.font.color.rgb=RGBColor.from_string(LAPIS)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(120); r=p.add_run("KANONISCHE ARBEITSFASSUNG · VERSION 1.0"); r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(GOLD)
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True,exist_ok=True)
    doc=Document(); styles(doc); setup(doc)
    doc.core_properties.title="Kana'anith – Sprach-, Namens- und Schriftbibel"; doc.core_properties.author="Weltbauprojekt Aleria"; doc.core_properties.keywords="Kana'anith, Tirnara, Sprache, Alphabet, Namen, Font"
    cover(doc)

    doc.add_heading("1. Sprachcharakter und Lautung",level=1)
    doc.add_paragraph("Kana’anith ist eine konsonantengetragene Sprache. Ein Wort besitzt einen bedeutungstragenden Konsonantenkern; die Vokale legen Aussprache, Wortart und grammatische Form fest. Die Schrift läuft von rechts nach links. In der praktischen lateinischen Umschrift werden Vokale vollständig notiert, im monumentalen Kana’anith hingegen meist nur durch Aleph, Ayin, Yodh und Waw angedeutet.")
    bullets(doc,[
        "Grundvokale: a, i und u; lange Formen ā, ī und ū. E und o sind jüngere Küsten- und Handelsvarianten.",
        "Bevorzugte Silben: CV, CVC, CVV und CVVC. Wörter beginnen gewöhnlich nicht mit zwei Konsonanten.",
        "Betonung: meist vorletzte Silbe; eine lange Endsilbe zieht die Betonung an sich.",
        "Harte Kultlaute: ḥ, ṭ, ṣ, q, ʾ und ʿ. In vereinfachter Umschrift sind HH, TT, SS, Q, A und AY zulässig.",
        "Drei Konsonanten bilden meist den eigentlichen Wortkern; zweikonsonantige Wurzeln werden in alten Amts- und Alltagswörtern bewahrt.",
    ])
    doc.add_heading("Vokalmuster",level=2)
    table(doc,["Muster","Funktion","Beispiel"],[
        ("CaCaC","greifbares Grundwort","Gaman: Reise"),("CiCiC","Beruf oder regelmäßige Tätigkeit","Gamir: Karawanenführer"),("CuCaC","erhöhter oder ritueller Begriff","Uralor: Gottheit"),("CaCāt","weibliche oder institutionelle Form","Rareshat: Königin"),("ma-CaC-im","geordnetes Kollektiv","Magamim: Karawane"),("CaC-um","Ort oder Gebiet","Merum: Brunnen"),
    ],[1.1,3.0,2.9],first_shaded=True)

    doc.add_heading("2. Die 22 Wurzelsilben",level=1)
    doc.add_paragraph("Die Zeichenbedeutungen sind bereits kanonisch. Für die Wortbildung erhält jedes Zeichen zusätzlich eine kurze Wurzelsilbe. Sie ist keine Umbenennung des Buchstabens, sondern seine produktive Sprachform.")
    table(doc,["Name","Laut","Wurzelsilbe","Bildbedeutung","Produktiver Bedeutungsraum"],[(name,ipa,root,picture,extended) for _,name,ipa,picture,extended,root in ALPHABET],[.85,.55,.85,1.5,3.25])
    doc.add_heading("Präfixe",level=2)
    table(doc,["Präfix","Bedeutung","Verwendung"],PREFIXES,[.8,1.8,4.4],first_shaded=True)
    doc.add_heading("Suffixe",level=2)
    table(doc,["Suffix","Bedeutung","Verwendung"],SUFFIXES,[.8,2.0,4.2],first_shaded=True)

    doc.add_page_break(); doc.add_heading("3. Namensbildung",level=1)
    doc.add_paragraph("Ein kana’anithischer Personenname verbindet zwei positive Zeichenbedeutungen. Der erste Namensstamm bezeichnet die erhoffte Grundanlage; der zweite nennt Aufgabe, Schutz oder Vermächtnis. Weibliche Traditionsformen verwenden weichere Endungen auf -a, ohne dass gewöhnliche Substantive ein grammatisches Geschlecht besitzen.")
    table(doc,["Baustein","Beispiele","Funktion"],[
        ("Erststamm","Aru-, Beti-, Gamu-, Meru-, Qenu-","gewünschte Anlage oder Herkunft"),("männlicher Zweitstamm","-bar, -gam, -dar, -zan, -yad, -lam, -mer, -sam, -resh, -tav","Aufgabe, Schutz oder öffentliches Vermächtnis"),("weiblicher Zweitstamm","-bara, -gama, -dara, -zana, -yada, -lama, -mera, -sama, -resha, -tava","traditionell weichere Namensform"),
    ],[1.25,2.8,3.75],first_shaded=True)

    doc.add_heading("100 männliche Namen",level=2)
    table(doc,["#","Name","Bausteine","Wörtlicher Sinn","Namensdeutung"],[(i,n,f"{first} + {end}",f"{m1} + {m2}",f"einer, der {m1.lower()} mit {m2.lower()} verbindet") for i,(n,first,end,m1,m2) in enumerate(MALE_NAMES,1)],[.35,1.05,1.25,2.25,2.1])
    doc.add_page_break(); doc.add_heading("100 weibliche Namen",level=2)
    table(doc,["#","Name","Bausteine","Wörtlicher Sinn","Namensdeutung"],[(i,n,f"{first} + {end}",f"{m1} + {m2}",f"eine, die {m1.lower()} mit {m2.lower()} verbindet") for i,(n,first,end,m1,m2) in enumerate(FEMALE_NAMES,1)],[.35,1.05,1.25,2.25,2.1])

    doc.add_page_break(); doc.add_heading("4. Grundwortschatz: 100 Wörter",level=1)
    doc.add_paragraph("Der Wortschatz folgt sichtbar den 22 Zeichenwurzeln. Die deutsche Übersetzung nennt die gebräuchliche Bedeutung; die wörtliche Herleitung zeigt die kana’anithische Denkweise.")
    groups={}
    for row in LEXICON: groups.setdefault(row[0],[]).append(row)
    num=1
    for category,rows in groups.items():
        doc.add_heading(category,level=2); body=[]
        for _,word,pos,derivation,literal,german in rows:
            body.append((num,word,pos,derivation,literal,german)); num+=1
        table(doc,["#","Kana’anith","Wortart","Herleitung","Wörtlich","Deutsch"],body,[.32,1.0,.72,1.25,1.7,2.01])

    doc.add_page_break(); doc.add_heading("5. Monumentalschrift und Font",level=1)
    doc.add_paragraph("Der Monumentalfont verwendet die 22 historischen RTL-Codepunkte U+10900 bis U+10915, zeichnet sie jedoch als eigenständigen tirnarischen Tempelschnitt. Die lateinische Tastaturbelegung dient der schnellen Eingabe; für korrektes Rechts-nach-links-Verhalten wandelt die mitgelieferte HTML-Demo die Umschrift in echte Kana’anith-Zeichen um.")
    doc.add_picture(str(CHART),width=Inches(7.0)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break(); doc.add_picture(str(SPECIMEN),width=Inches(7.0)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Eingaberegeln",level=2)
    table(doc,["Eingabe","Ergebnis","Hinweis"],[
        ("SH","Shin, š","Zahnlaut"),("HH","Heth, ḥ","harter Kehl- und Mauerlaut"),("TT","Teth, ṭ","emphatisches t"),("SS","Sade, ṣ","emphatisches s"),("U+10900–U+10915","echte Kana’anith-Zeichen","automatische RTL-Laufrichtung"),("A–Z","praktische Aliasbelegung","für Desktopprogramme ohne Transliterator"),
    ],[1.2,2.0,3.8],first_shaded=True)

    doc.add_heading("6. Bestandsprüfung",level=1)
    counts=Counter(row[0] for row in LEXICON)
    table(doc,["Bestand","Anzahl","Status"],[
        ("Konsonantenzeichen",len(ALPHABET),"22 kanonische Zeichen und Wurzelsilben"),("Präfixe",len(PREFIXES),"14 produktive Ableitungen"),("Suffixe",len(SUFFIXES),"18 produktive Endungen"),("Männliche Namen",len(MALE_NAMES),"100 eindeutige Namen"),("Weibliche Namen",len(FEMALE_NAMES),"100 eindeutige Namen"),("Kulturwörter",len(LEXICON),"100 eindeutige Wörter in sieben Themenfeldern"),
    ],[2.2,.8,4.0],first_shaded=True)
    trailing = doc.paragraphs[-1]
    trailing._element.getparent().remove(trailing._element)
    doc.save(OUT); print(OUT)


if __name__=="__main__": build()
